import os
import json
import logging
from dotenv import load_dotenv
from marketing_ai_agent import MarketingAIAgent
from sales_ai_agent import SalesAIAgent
from strategy_ai_agent import StrategyAIAgent
from analytics_ai_agent import AnalyticsAIAgent
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from colorama import init, Fore, Style

# Initialize colorama
init(autoreset=True)

# Load environment variables
load_dotenv()

# Set up logging
logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')

def create_styled_document(content, language='english'):
    doc = Document()
    
    # Create and apply styles
    title_style = doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = 'Arial'
    title_style.font.size = Pt(24)
    title_style.font.color.rgb = RGBColor(0, 112, 192)  # Blue color
    title_style.font.bold = True
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    heading_style = doc.styles.add_style('CustomHeading', WD_STYLE_TYPE.PARAGRAPH)
    heading_style.font.name = 'Calibri'
    heading_style.font.size = Pt(16)
    heading_style.font.color.rgb = RGBColor(46, 116, 181)  # Dark blue color
    heading_style.font.bold = True
    
    body_style = doc.styles.add_style('CustomBody', WD_STYLE_TYPE.PARAGRAPH)
    body_style.font.name = 'Georgia'
    body_style.font.size = Pt(11)
    body_style.paragraph_format.space_after = Pt(12)
    
    # Add content to the document
    title = "Marketing Team Discussion" if language == 'english' else "Marketing-Team-Diskussion"
    doc.add_paragraph(title, style='CustomTitle')
    
    for section in content:
        doc.add_paragraph(section['title'], style='CustomHeading')
        doc.add_paragraph(section['content'], style='CustomBody')
    
    # Add page numbers
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = ("Page " if language == 'english' else "Seite ") + "{ PAGE }"
    footer_para.style = doc.styles['Footer']
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save the document
    filename = 'marketing_plan.docx' if language == 'english' else 'marketingplan.docx'
    doc.save(filename)
    print(f"Marketing plan saved as '{filename}'")

def check_api_key():
    api_key = os.getenv("OPENROUTER_API_KEY")
    print(f"Current API key value: {api_key}")
    if not api_key or api_key.strip() == "":
        print("OpenRouter API key not found or empty in environment variables.")
        new_key = input("Please enter your OpenRouter API key: ").strip()
        if new_key:
            with open(".env", "w") as env_file:
                env_file.write(f"OPENROUTER_API_KEY={new_key}")
            print("API key has been added to .env file.")
            os.environ["OPENROUTER_API_KEY"] = new_key
            api_key = new_key
        else:
            print("No API key provided. Exiting.")
            exit(1)
    return api_key

class MarketingTeam:
    def __init__(self):
        self.marketing_agent = MarketingAIAgent()
        self.sales_agent = SalesAIAgent()
        self.strategy_agent = StrategyAIAgent()
        self.analytics_agent = AnalyticsAIAgent()

    def discuss_marketing_plan(self, product, additional_info=None):
        print(f"{Fore.CYAN}Marketing Team discussing: {product}{Style.RESET_ALL}\n", flush=True)
        logging.info(f"Starting marketing plan discussion for {product}")

        content = []
        language = additional_info.get('language', 'english')

        # Marketing Agent's initial input
        print(f"{Fore.RED}Marketing Agent: ", end='', flush=True)
        marketing_input = self.marketing_agent.generate_campaign_idea(product, additional_info=additional_info)
        logging.debug(f"Marketing Agent response: {marketing_input}")
        content.append({"title": "Initial Marketing Campaign Idea", "content": marketing_input})
        print(f"{Style.RESET_ALL}\n", flush=True)

        # Sales Agent's response to Marketing
        print(f"{Fore.GREEN}Sales Agent: ", end='', flush=True)
        sales_input = self.sales_agent.respond_to_agent(marketing_input, language=language)
        logging.debug(f"Sales Agent response: {sales_input}")
        content.append({"title": "Sales Agent Feedback", "content": sales_input})
        print(f"{Style.RESET_ALL}\n", flush=True)

        # Strategy Agent's input based on Marketing and Sales
        print(f"{Fore.YELLOW}Strategy Agent: ", end='', flush=True)
        strategy_input = self.strategy_agent.analyze_market_trends(product, f"{marketing_input}\n{sales_input}", language=language)
        logging.debug(f"Strategy Agent response: {strategy_input}")
        content.append({"title": "Market Trends Analysis", "content": strategy_input})
        print(f"{Style.RESET_ALL}\n", flush=True)

        # Analytics Agent's input based on all previous inputs
        print(f"{Fore.MAGENTA}Analytics Agent: ", end='', flush=True)
        analytics_input = self.analytics_agent.analyze_target_audience(product, f"{marketing_input}\n{sales_input}\n{strategy_input}", language=language)
        logging.debug(f"Analytics Agent response: {analytics_input}")
        content.append({"title": "Target Audience Analysis", "content": analytics_input})
        print(f"{Style.RESET_ALL}\n", flush=True)

        # Marketing Agent's final input based on all feedback
        print(f"{Fore.RED}Marketing Agent (Final): ", end='', flush=True)
        final_marketing_input = self.marketing_agent.generate_campaign_idea(product, additional_info=f"{sales_input}\n{strategy_input}\n{analytics_input}")
        logging.debug(f"Final Marketing Agent response: {final_marketing_input}")
        content.append({"title": "Final Marketing Campaign Idea", "content": final_marketing_input})
        print(f"{Style.RESET_ALL}\n", flush=True)

        # Final plan synthesis
        print(f"{Fore.BLUE}Synthesizing final plan...\n", flush=True)
        final_plan = self.synthesize_plan(final_marketing_input, sales_input, strategy_input, analytics_input, product, additional_info)
        logging.debug(f"Final plan: {final_plan}")
        content.append({"title": "Final Marketing Plan", "content": final_plan})

        # Create styled Word document
        create_styled_document(content, language)
        logging.info("Marketing plan discussion completed")

    def synthesize_plan(self, marketing, sales, strategy, analytics, product, additional_info):
        language = additional_info.get('language', 'english')
        prompt = f"""
        Synthesize a comprehensive marketing plan for {product} based on the following inputs:
        Marketing: {marketing}
        Sales: {sales}
        Strategy: {strategy}
        Analytics: {analytics}

        Additional Information:
        Target Audience: {additional_info.get('target_audience', 'Not specified')}
        Marketing Goals: {additional_info.get('marketing_goals', 'Not specified')}
        Budget: {additional_info.get('budget', 'Not specified')}

        Provide a cohesive plan that incorporates insights from all agents, specifically tailored for {product},
        taking into account the additional information provided.
        """
        return self.marketing_agent.call_openrouter_api(prompt, language=language)

def main():
    api_key = check_api_key()
    os.environ["OPENROUTER_API_KEY"] = api_key
    team = MarketingTeam()
    
    print("Please provide details for your marketing plan:")
    print("Enter product name and any additional important information (e.g., target audience, goals, budget, competitors, unique selling points, timeframe, etc.):")
    print("Press Enter twice when you're done.")
    
    brief_lines = []
    while True:
        line = input()
        if line.strip() == "":
            if brief_lines:  # If we have content and encounter an empty line, break
                break
        else:
            brief_lines.append(line)
    
    brief = " ".join(brief_lines)
    
    if not brief:
        print("Error: Please provide information about the product and marketing context. Please try again.")
        return
    
    language = ''
    while language not in ['english', 'german']:
        language = input("Enter preferred language (english/german): ").strip().lower()
        if language not in ['english', 'german']:
            print("Invalid language. Please enter either 'english' or 'german'.")
    
    # Extract product name from the brief (assuming it's the first sentence up to the first period)
    product = brief.split('.')[0].strip()
    additional_info_dict = {
        "additional_info": brief,
        "language": language
    }
    
    team.discuss_marketing_plan(product, additional_info_dict)
    
    print(f"Thank you for using the Marketing Team AI. Your marketing plan has been generated and saved in {language}.")

if __name__ == "__main__":
    main()
